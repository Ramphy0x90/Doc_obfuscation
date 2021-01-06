import sys
import docx
from termcolor import colored

banner = colored('''
    ____                                        __           __    ____                      __  _           
   / __ \\____  _______  ______ ___  ___  ____  / /_   ____  / /_  / __/_  ________________ _/ /_(_)___  ____ 
  / / / / __ \\/ ___/ / / / __ `__ \\/ _ \\/ __ \\/ __/  / __ \\/ __ \\/ /_/ / / / ___/ ___/ __ `/ __/ / __ \\/ __ \\
 / /_/ / /_/ / /__/ /_/ / / / / / /  __/ / / / /_   / /_/ / /_/ / __/ /_/ (__  ) /__/ /_/ / /_/ / /_/ / / / /
/_____/\\____/\\___/\\__,_/_/ /_/ /_/\\___/_/ /_/\\__/   \\____/_.___/_/  \\__,_/____/\\___/\\__,_/\\__/_/\\____/_/ /_/ 

                                                                                    by {}
''', 'yellow').format(colored('Ramphy Aquino Nova', 'blue'))

refer_red = colored('[*]', 'red')
refer_yellow = colored('[*]', 'yellow')


def getBanner() -> str:
    text = '''                                                                                                            
    {}
    
{} Usage: main.py (-t, -w) [file_path]

{} Examples: main.py -t 'hello world'
              main.py -w my_document.docx
    '''.format(banner, refer_red, refer_red)

    return text


def checkArguments() -> list:
    args_ref = ['-t', '-w', '-p']
    params = []

    for arg in args_ref:
        if arg in sys.argv:
            params.append(arg)
            if len(sys.argv) == 3:
                params.append(sys.argv[-1])

    return params


def obfuscateText(text: str) -> list:
    main_chars = {
        'A': '\u0391',  # Greek chars
        'B': '\u0392',
        'C': '\u03f9',
        'H': '\u0397',
        'T': '\u03a4',
        'P': '\u0420',  # Cyrillic chars
        'M': '\u041c',
        'a': '\u0430',
        'c': '\u0441',
        's': '\u0455',
        'i': '\u0456',
        'x': '\u0445',
        'n': '\u0578',  # Armenian chars
        'h': '\u0570'
    }

    final_text = ''
    changed_chars = 0
    original_ascii = []
    obfuscated_ascii = []

    for char in text:
        original_ascii.append(ord(char))

        try:
            obfuscated_ascii.append(ord(main_chars[char]))
            final_text += main_chars[char]
            changed_chars += 1
        except KeyError:
            obfuscated_ascii.append(ord(char))
            final_text += char

    return [final_text, str(original_ascii)[1:-1], str(obfuscated_ascii)[1:-1], changed_chars]


if __name__ == '__main__':
    userArguments = checkArguments()

    if len(userArguments) != 2:
        print(getBanner())
    else:
        print(banner)

        if userArguments[0] == '-t':
            result = obfuscateText(userArguments[1])

            print('{} Original ascii values: {}'.format(refer_red, colored(result[1], 'blue')))
            print('{} Obfuscated ascii values: {}'.format(refer_red, colored(result[2], 'blue')))
            print('{} Obfuscated text: {}'.format(refer_red, colored(result[0], 'yellow')))

        elif userArguments[0] == '-w':
            try:
                new_document_name = 'obfuscated-{}'.format(userArguments[1])

                document1 = docx.Document(userArguments[1])
                document1.save(new_document_name)

                document2 = docx.Document(new_document_name)

                print('{} Document name: {}'.format(refer_red, colored(userArguments[1], 'blue')))
                print('{} Paragraph number: {}'.format(refer_red, colored(len(document2.paragraphs), 'blue')))

                total_changed_chars = 0

                for paragraph in document2.paragraphs:
                    result = obfuscateText(paragraph.text)
                    paragraph.text = result[0]
                    total_changed_chars += result[-1]

                document2.save(new_document_name)

                print('{} Letters replaced: {}'.format(refer_red, colored(total_changed_chars, 'blue')))
                print('{} New obfuscated file: {}'.format(refer_red, colored(new_document_name, 'blue')))

            except docx.opc.exceptions.PackageNotFoundError:
                print('\n[!] Document not found')
